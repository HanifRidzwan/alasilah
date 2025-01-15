from flask import Flask, render_template, request, redirect, flash, session, send_file
from argon2 import PasswordHasher, exceptions
from argon2.exceptions import VerifyMismatchError
from models.question_generator import generate_questions, save_questions_to_docx
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import sqlite3
import os
import smtplib
import random
import re


app = Flask(__name__)
app.secret_key = 'your_secret_key_here'

ph = PasswordHasher()

def get_db_connection():
    return sqlite3.connect('users.db')

# Ensure export directory exists
EXPORT_DIR = 'exports'
os.makedirs(EXPORT_DIR, exist_ok=True)

@app.route('/')
def home():
    print("Rendering the home page.")
    return render_template('index.html')

# SMTP setup
SMTP_SERVER = "smtp.mail.yahoo.com"
SMTP_PORT = 587
EMAIL_ADDRESS = "dragunov_molotovhr@yahoo.com.my"
EMAIL_PASSWORD = "zeokilktrgwfjfjf"  # Replace with your Yahoo app password


def get_db_connection():
    return sqlite3.connect('users.db')

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def send_verification_email(email, code):
    try:
        # Create a MIME multipart message
        msg = MIMEMultipart()
        msg['Subject'] = "Al-As'ilah - Your Verification Code"
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = email

        # Email content with HTML for better styling
        body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <h2 style="color: #6c63ff;">Al-As'ilah - Verification Code</h2>
            <p>Dear User,</p>
            <p>Thank you for signing up for <strong>Al-As'ilah</strong>. Please use the verification code below to complete your registration:</p>
            <div style="padding: 10px; background-color: #f3f4f6; border-radius: 5px; display: inline-block; font-size: 20px; font-weight: bold; color: #6c63ff;">
                {code}
            </div>
            <p>This code is valid for <strong>1 minute</strong>. If you did not request this code, please ignore this email.</p>
            <br>
            <p>Thank you,</p>
            <p><strong>The Al-As'ilah Team</strong></p>
            <hr style="border: none; border-top: 1px solid #ddd;">
            <p style="font-size: 12px; color: #666;">This email was sent from a no-reply address. Please do not reply to this email.</p>
        </body>
        </html>
        """
        # Attach the HTML content
        msg.attach(MIMEText(body, 'html'))

        # SMTP setup and sending email
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)  # Use your existing email credentials
            server.send_message(msg)

        print(f"Verification email sent to {email}.")
    except Exception as e:
        print(f"Failed to send email: {e}")



@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        email = request.form['email']
        staff_id = request.form['staff_id']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        print("Signup request received:", email, staff_id)

        # Check if passwords match
        if password != confirm_password:
            flash('Passwords do not match!', 'error')
            print("Password mismatch.")
            return redirect('/signup')

        # Validate password strength
        if not re.fullmatch(r'^(?=.*[A-Z])(?=.*[0-9])(?=.*[!@#$%^&*])[A-Za-z\d!@#$%^&*]{8,}$', password):
            flash('Password must be at least 8 characters long and include uppercase letters, numbers, and special symbols.', 'error')
            print("Weak password.")
            return redirect('/signup')

        try:
            # Check if email already exists in the database
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('SELECT * FROM users WHERE email = ?', (email,))
                if cursor.fetchone():
                    flash('An account with this email already exists!', 'error')
                    print("Email already exists in the database.")
                    return redirect('/signup')

                # Check if staff ID already exists
                cursor.execute('SELECT * FROM users WHERE staff_id = ?', (staff_id,))
                if cursor.fetchone():
                    flash('This Staff ID is already registered!', 'error')
                    print("Staff ID already exists in the database.")
                    return redirect('/signup')

                # Generate verification code
                code = ''.join(random.choices('0123456789', k=6))
                expiration_time = datetime.now() + timedelta(minutes=1)

                print("Generated verification code:", code)

                # Store verification code in the database
                cursor.execute('''
                    INSERT INTO verification_codes (email, code, expiration_time) 
                    VALUES (?, ?, ?)
                ''', (email, code, expiration_time))
                conn.commit()

                print("Verification code stored in the database.")

            # Send email with the verification code
            send_verification_email(email, code)

            # Store user details temporarily in the session
            session['signup_details'] = {
                'email': email,
                'staff_id': staff_id,
                'password': ph.hash(password)  # Hash password securely
            }

            print("Session data set:", session['signup_details'])

            flash('Verification code sent to your email. Please verify.', 'success')
            return redirect('/verify')

        except sqlite3.Error as db_error:
            flash('A database error occurred. Please try again.', 'error')
            print(f"Database error during sign-up: {db_error}")

        except Exception as e:
            flash('An unexpected error occurred during sign-up. Please try again.', 'error')
            print(f"Unexpected error during sign-up: {e}")

    return render_template('signup.html')


@app.route('/verify', methods=['GET', 'POST'])
def verify():
    if request.method == 'POST':
        code = request.form['code']
        email = session.get('signup_details', {}).get('email')

        print("Verification request received for email:", email)
        print("Code entered by user:", code)

        if not email:
            flash('Session expired. Please sign up again.', 'error')
            print("Session expired or email not found in session.")
            return redirect('/signup')

        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT code, expiration_time FROM verification_codes 
                    WHERE email = ?
                ''', (email,))
                record = cursor.fetchone()

                print("Database record fetched for verification:", record)

                if not record:
                    flash('Verification code not found.', 'error')
                    print("No code found for this email.")
                    return redirect('/verify')

                db_code, expiration_time = record
                print("Code in database:", db_code)
                print("Code expiration time:", expiration_time)
                print("Current time:", datetime.now())

                # Check code and expiration time
                if db_code != code:
                    flash('Invalid verification code.', 'error')
                    print("Verification code does not match.")
                    return redirect('/verify')

                if datetime.strptime(expiration_time, '%Y-%m-%d %H:%M:%S.%f') < datetime.now():
                    flash('Verification code has expired.', 'error')
                    print("Verification code expired.")
                    return redirect('/verify')

                # Save user to the database
                details = session.pop('signup_details')
                cursor.execute('''
                    INSERT INTO users (email, staff_id, password) 
                    VALUES (?, ?, ?)
                ''', (details['email'], details['staff_id'], details['password']))
                conn.commit()

                print("User successfully verified and added to database.")
                flash('Account verified and created successfully!', 'success')
                return redirect('/login')

        except Exception as e:
            flash('An error occurred during verification. Please try again.', 'error')
            print(f"Error during verification: {e}")

    return render_template('verify.html')

# resend route
@app.route('/resend', methods=['GET'])
def resend_code():
    # Check if session contains signup details
    signup_details = session.get('signup_details')
    if not signup_details:
        flash('Session expired. Please sign up again.', 'error')
        return redirect('/signup')

    email = signup_details.get('email')

    try:
        # Generate a new verification code
        code = ''.join(random.choices('0123456789', k=6))
        expiration_time = datetime.now() + timedelta(minutes=1)

        print("Generated new verification code:", code)

        # Update the verification code in the database
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE verification_codes
                SET code = ?, expiration_time = ?
                WHERE email = ?
            ''', (code, expiration_time, email))
            conn.commit()

        # Send the new verification code via email
        send_verification_email(email, code)

        flash('A new verification code has been sent to your email.', 'success')
        return redirect('/verify')

    except sqlite3.Error as db_error:
        flash('A database error occurred. Please try again.', 'error')
        print(f"Database error during resend: {db_error}")

    except Exception as e:
        flash('An error occurred while resending the verification code. Please try again.', 'error')
        print(f"Error during resend: {e}")

    return redirect('/verify')


# login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('SELECT id, password FROM users WHERE email = ?', (email,))
                user = cursor.fetchone()

                if user:
                    try:
                        # Verify the hashed password using Argon2
                        ph.verify(user[1], password)
                        session['user_id'] = user[0]  # Store user ID in session
                        flash('Login successful!', 'success')
                        return redirect('/dashboard')  # Redirect to dashboard
                    except VerifyMismatchError:
                        flash('Invalid password!', 'error')  # Handle incorrect password
                else:
                    flash('Email not found!', 'error')  # Handle case where email does not exist
        except sqlite3.Error as e:
            flash('Database error occurred. Please try again later.', 'error')
            print(f"Database error: {e}")  # Log the error for debugging
        except Exception as e:
            flash('An unexpected error occurred. Please try again.', 'error')
            print(f"Unexpected error: {e}")  # Log unexpected errors

    return render_template('login.html')



# profile

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'user_id' not in session:  # Ensure user is logged in
        return redirect('/login')

    user_id = session['user_id']
    conn = get_db_connection()
    cursor = conn.cursor()

    if request.method == 'POST':
        email = request.form['email']
        staff_id = request.form['staff_id']
        current_password = request.form['current_password']
        new_password = request.form['new_password']
        confirm_new_password = request.form['confirm_new_password']

        # Fetch current user details
        cursor.execute('SELECT email, staff_id, password FROM users WHERE id = ?', (user_id,))
        user = cursor.fetchone()

        # Validate current password
        if not ph.verify(user[2], current_password):
            flash('Incorrect current password!', 'error')
            return redirect('/profile')

        # Update user details
        if new_password and new_password == confirm_new_password:
            hashed_new_password = ph.hash(new_password)
            cursor.execute('UPDATE users SET email = ?, staff_id = ?, password = ? WHERE id = ?',
                           (email, staff_id, hashed_new_password, user_id))
        else:
            cursor.execute('UPDATE users SET email = ?, staff_id = ? WHERE id = ?',
                           (email, staff_id, user_id))

        conn.commit()
        conn.close()
        flash('Profile updated successfully!', 'success')
        return redirect('/profile')

    # Pre-fill form with existing user data
    cursor.execute('SELECT email, staff_id FROM users WHERE id = ?', (user_id,))
    user = cursor.fetchone()
    conn.close()

    return render_template('profile.html', email=user[0], staff_id=user[1])


    # dashboard
import re

@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if 'user_id' not in session:
        flash('You need to log in to access the dashboard.', 'error')
        return redirect('/login')

    if request.method == 'POST':
        query = request.form['query']
        question_count = int(request.form['question_count'])
        question_type = request.form['question_type']

        try:
            # Generate questions using the model
            generated_questions = generate_questions(query, question_count, question_type)

            # Ensure Answer Options are properly split
            for question in generated_questions:
                if "Answer Options" in question and isinstance(question["Answer Options"], str):
                    question["Answer Options"] = [opt.strip() for opt in question["Answer Options"].split("؛")]

            # Save query and questions to session
            if 'queries' not in session:
                session['queries'] = []
            session['queries'].append({
                'query': query,
                'questions': generated_questions
            })
            session.modified = True

            flash(f"Generated {question_count} questions for '{query}'.", 'success')
        except Exception as e:
            flash(f"Error generating questions: {e}", 'error')

    queries = session.get('queries', [])
    return render_template('dashboard.html', queries=queries)



from flask import send_file, jsonify, request
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

@app.route('/export_question', methods=['POST'])
def export_question():
    """Export filtered questions based on user query to a .docx file."""
    try:
        # Fetch the filtered questions from the request
        questions = request.json.get('questions', [])
        print("Received questions for export:", questions)  # Debugging log

        if not questions:
            return jsonify({"error": "No questions to export"}), 400

        # Create a new Word document
        doc = Document()

        # Set global RTL alignment
        for section in doc.sections:
            section.right_to_left = True

        # Set font style for Arabic text
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        # Add a title to the document
        title = doc.add_heading('الأسئلة المولدة', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add each question with its options
        for question_data in questions:
            parts = question_data.split("\n")
            question_text = parts[0].strip()  # The first line is always the question text
            options = [opt.strip() for opt in parts[1:] if opt.strip()]  # Remaining lines are options (if any)

            # Add question text with explicit RTL enforcement
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(question_text)
            question_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            question_para.paragraph_format.right_to_left = True

            # Explicitly set RTL for content in run
            rtl_element = question_run._element.get_or_add_rPr().get_or_add_rtl()
            rtl_element.text = "true"

            if options:
                # Add options for MCQ questions
                for option in options:
                    option_para = doc.add_paragraph()
                    option_run = option_para.add_run(option)
                    option_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    option_para.paragraph_format.right_to_left = True
                    rtl_element = option_run._element.get_or_add_rPr().get_or_add_rtl()
                    rtl_element.text = "true"
            else:
                # Add placeholder for short-answer questions
                placeholder_para = doc.add_paragraph()
                placeholder_para.add_run("__________").bold = True
                placeholder_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                placeholder_para.paragraph_format.right_to_left = True

            # Add spacing between questions
            doc.add_paragraph()

        # Save the document to memory
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Return the document as a downloadable file
        return send_file(doc_io, as_attachment=True, download_name="questions.docx")

    except Exception as e:
        print("Error during export:", e)
        return jsonify({"error": "An error occurred during export"}), 500





from flask import session, redirect

@app.route('/logout')
def logout():
    session.clear()  # Clears all session data
    return redirect('/')  # Redirect to homepage



if __name__ == '__main__':
    print("Starting the Flask application.")
    app.run(debug=True)
