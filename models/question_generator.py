import pandas as pd 
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from camel_tools.tokenizers.word import simple_word_tokenize
from camel_tools.utils.dediac import dediac_ar
from io import BytesIO
import csv


# Load the corpus
try:
    corpus = pd.read_csv("data/compile.csv")  # Adjust path to align with project structure
except FileNotFoundError:
    raise Exception("Corpus file not found. Ensure compile.csv is in the data folder.")

# Define Arabic stopwords
arabic_stopwords = ["في", "على", "من", "إلى", "و", "ما", "هو", "هل", "لم", "هذا"]

# Training data with expanded example queries
file_path = "data/training_data.csv"
df = pd.read_csv(file_path, encoding="utf-8-sig")

# Clean the 'query' column: remove extra single quotes
df["query"] = df["query"].str.strip().str.replace("'", "")

# Convert to list of tuples
training_data = list(df.itertuples(index=False, name=None))

# Print cleaned data
print(training_data[:5])

# Separate inputs and intents for training
train_texts, train_labels = zip(*training_data)

# Vectorize input text using TF-IDF
vectorizer = TfidfVectorizer()
X_train = vectorizer.fit_transform(train_texts)

# Train the Naive Bayes classifier
classifier = MultinomialNB()
classifier.fit(X_train, train_labels)

def classify_intent(user_input):
    """Classify the intent of the user's query."""
    X_test = vectorizer.transform([user_input])
    predicted_intent = classifier.predict(X_test)[0]
    return predicted_intent

def preprocess_text(text):
    """Preprocess Arabic text by removing diacritics, tokenizing, and removing stop words."""
    normalized_text = dediac_ar(text)
    tokens = simple_word_tokenize(normalized_text)
    filtered_tokens = [token for token in tokens if token not in arabic_stopwords]
    return " ".join(filtered_tokens)

def filter_questions(intent, question_type, num_questions):
    """
    Filter questions based on identified intent and question type.
    Returns a sample of questions as a DataFrame.
    """
    # Filter corpus based on question type and intent
    filtered = corpus[
        (corpus["Question Type"] == question_type) & 
        (corpus["Topic"].str.contains(intent, na=False))
    ]
    
    # Adjust sample size if fewer questions are available
    return filtered.sample(min(num_questions, len(filtered))) if not filtered.empty else pd.DataFrame()

def generate_questions(user_input, num_questions, question_type):
    """
    Generate questions based on the user's query, desired number of questions, and question type.
    Returns questions as a list of dictionaries suitable for JSON serialization.
    """
    # Classify the user's query to get the intent
    intent = classify_intent(user_input)

    # Generate questions based on identified intent
    questions_df = filter_questions(intent, question_type, num_questions)

    # Check if any questions were found
    if questions_df.empty:
        return [{"Question Text": "No questions found for this query."}]

    # Convert the DataFrame to a list of dictionaries for easier JSON serialization
    questions = []
    for _, row in questions_df.iterrows():
        question = {
            "Question Text": row["Question Text"],
            "Answer Options": row["Answer Options"] if pd.notna(row["Answer Options"]) else None,
            "Correct Answer": row["Correct Answer"] if pd.notna(row["Correct Answer"]) else None,
            "Paragraph": row["Paragraph"] if pd.notna(row["Paragraph"]) else None,
        }
        questions.append(question)

    return questions

def save_questions_to_docx(questions, filename='generated_questions.docx'):
    """
    Save the questions in a professional Arabic exam format.
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt

    doc = Document()

    # Set the document title
    title = doc.add_heading(level=1)
    run = title.add_run("Generated Arabic Exam Questions")
    run.font.name = "Traditional Arabic"
    run.font.size = Pt(16)
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add each question
    for idx, question in enumerate(questions, 1):
        # Add the question text
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = para.add_run(f"{idx}. {question.get('Question Text', 'No question text available')}")
        run.font.name = "Traditional Arabic"
        run.font.size = Pt(14)

        # Add the answer options
        if question.get("Answer Options"):
            options = question["Answer Options"].split(", ")
            for option_label, option_text in zip(["أ.", "ب.", "ج.", "د."], options):
                option_para = doc.add_paragraph()
                option_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run = option_para.add_run(f"{option_label} {option_text}")
                run.font.name = "Traditional Arabic"
                run.font.size = Pt(12)

        # Add spacing after the question
        doc.add_paragraph()

    # Save the document
    if isinstance(filename, BytesIO):  # Support in-memory saving for Flask
        doc.save(filename)
    else:
        doc.save(filename)
